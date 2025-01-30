import streamlit as st
import google.generativeai as genai
from PIL import Image
import db_manager as db
from docx import Document
from io import BytesIO

genai.configure(api_key=st.secrets["API_KEY"])

model = genai.GenerativeModel(model_name="gemini-1.5-flash-002")

if "started" not in st.session_state:
    st.session_state["analyze_image_finished"] = False
    st.session_state["chat_control"] = []
    st.session_state.response_download = None
    st.session_state.uploaded_files = None
    st.session_state.started = False

st.title("EtudIAnt : üìù Cr√©ateur de contr√¥les")

if st.session_state.started == False:
    level = st.selectbox('S√©lectionne ton niveau : ', ["6√®me","5√®me","4√®me","3√®me","Seconde","Premiere","Terminale"])
    subject = st.selectbox('S√©lectionne la mati√®re de ta fiche de r√©vision:', ["Fran√ßais", "Math√©matiques", "Histoire-G√©ographie-EMC", "Sciences et Vie de la Terre", "Physique Chimie", "Anglais","Allemand", "Espagnol"])   
    difficulty = st.slider("D√©finis la difficult√©e du contr√¥le de 1 √† 10 :", 1, 10)
    prompt = f"Voici un groupe d'images d'un cours de niveau {level}. Cr√©e un contr√¥le de difficult√© : {difficulty} sur 10. Comme au college ou lyc√©e dessus en adaptant la difficult√©e en fonction du niveau er de la mati√®re : {subject}. R√©pond en parlant francais, jamais en anglais. Ne fais pas ton introduction dans la r√©ponse, fais directement le contr√¥le"
    st.session_state.uploaded_files = st.file_uploader("T√©l√©charge les photos de tes cours.", type=["png", "jpg", "jpeg", "bmp"], accept_multiple_files=True)

def display_images(files):
    images = []
    for file in files:
        image_pil = Image.open(file)
        image_pil.resize((256, 256))
        st.image(image_pil, caption=file.name, use_container_width=True)
        images.append(image_pil)
    return images

def create_dynamic_word_doc(response):
    doc = Document()
    doc.add_heading('R√©ponse de l\'IA', level=1)

    lines = response.split("\n")
    table = None
    
    for line in lines:
        if "|" in line:
            cells = [cell.strip() for cell in line.split("|") if cell.strip()]
            if not table:
                table = doc.add_table(rows=1, cols=len(cells))
                hdr_cells = table.rows[0].cells
                for i, cell in enumerate(cells):
                    hdr_cells[i].text = cell
            else:
                row_cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    row_cells[i].text = cell
        else:
            doc.add_paragraph(line)

    return doc


place_holder_button = st.empty()
can_button = st.session_state.started
if st.session_state.uploaded_files and len(st.session_state.uploaded_files) > 0 or st.session_state.started == True and st.session_state.chat_control != []:
    if place_holder_button.button("Cr√©er un contr√¥le sur ce cours",disabled=can_button):
        if db.can_user_make_request():
            images = display_images(st.session_state.uploaded_files)
            if not st.session_state["analyze_image_finished"]:
                with st.spinner("L'EtudIAnt reflechit..."):
                    place_holder_button.empty()
                    response = model.generate_content([prompt]+ images)
                st.session_state.response_download = response.text
                st.session_state["chat_control"].append({"role": "assistant", "content": response.text})
                st.session_state.started = True
                db.consume_request()
                st.rerun()
        else:
            st.error("Votre quotas de requ√™tes par jour est termin√©, revenez demain pour utiliser l'EtudIAnt.")

if st.session_state.started == True:
    doc = create_dynamic_word_doc(st.session_state.response_download)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button(
    label="T√©l√©charger la r√©ponse en Word (b√™ta)",
    data=buffer,
    file_name="response.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if "analyze_image_finished" in st.session_state:
    for message in st.session_state["chat_control"]:
        if message["role"] == "assistant":
            with st.chat_message("assistant"):
                st.write(f"**IA** : {message['content']}")
                st.write("Si tu veux gagner des √©toiles ‚≠ê r√©soud ce contr√¥le et envoie le moi !")

